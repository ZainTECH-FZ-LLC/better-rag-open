"""DOCX generator — creates Word documents from structured specs."""

from __future__ import annotations

import asyncio
import uuid
from pathlib import Path

import structlog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.document_generation.base import DocumentGenerator, DocumentSpec, GeneratedDocument

logger = structlog.get_logger()


class DOCXGenerator(DocumentGenerator):
    """
    Generates Word documents using python-docx.

    Section types:
    - heading: Section heading with level
    - paragraph: Body text
    - bullets: Bullet point list
    - table: Data table
    - numbered_list: Ordered list
    """

    async def generate(self, spec: DocumentSpec, output_dir: Path) -> GeneratedDocument:
        return await asyncio.to_thread(self._generate_sync, spec, output_dir)

    def _generate_sync(self, spec: DocumentSpec, output_dir: Path) -> GeneratedDocument:
        doc = Document()

        # Document title
        title_para = doc.add_heading(spec.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle / metadata line
        subtitle = spec.data.get("subtitle", "")
        if subtitle:
            sub_para = doc.add_paragraph(subtitle)
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.runs[0].font.size = Pt(12)
            sub_para.runs[0].font.italic = True

        doc.add_paragraph("")  # Spacer

        # Build sections
        for section in spec.sections:
            section_type = section.get("type", "paragraph")

            if section_type == "heading":
                level = section.get("level", 1)
                doc.add_heading(section.get("title", ""), level=level)

            elif section_type == "paragraph":
                text = section.get("text", "")
                if text:
                    para = doc.add_paragraph(text)
                    para.style.font.size = Pt(11)

            elif section_type == "bullets":
                heading = section.get("title")
                if heading:
                    doc.add_heading(heading, level=2)
                for bullet in section.get("items", []):
                    doc.add_paragraph(bullet, style="List Bullet")

            elif section_type == "numbered_list":
                heading = section.get("title")
                if heading:
                    doc.add_heading(heading, level=2)
                for item in section.get("items", []):
                    doc.add_paragraph(item, style="List Number")

            elif section_type == "table":
                heading = section.get("title")
                if heading:
                    doc.add_heading(heading, level=2)

                headers = section.get("headers", [])
                rows_data = section.get("rows", [])

                if headers and rows_data:
                    table = doc.add_table(
                        rows=len(rows_data) + 1,
                        cols=len(headers),
                        style="Table Grid",
                    )

                    # Header row
                    for j, header in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = str(header)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)

                    # Data rows
                    for i, row in enumerate(rows_data):
                        for j, value in enumerate(row):
                            if j < len(headers):
                                cell = table.rows[i + 1].cells[j]
                                cell.text = str(value)
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.font.size = Pt(10)

                doc.add_paragraph("")  # Spacer after table

            elif section_type == "chart":
                heading = section.get("title")
                if heading:
                    doc.add_heading(heading, level=2)

                chart_spec = section.get("chart", {})
                self._add_chart(doc, chart_spec, output_dir)

        # Save
        filename = f"{_sanitize(spec.title)}_{uuid.uuid4().hex[:8]}.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))

        size = filepath.stat().st_size
        logger.info("docx_generator.saved", filename=filename, size=size)

        return GeneratedDocument(
            filename=filename,
            filepath=filepath,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_type="docx",
            size_bytes=size,
        )

    def _add_chart(self, doc: Document, chart_spec: dict, output_dir: Path) -> None:
        """Render a chart via matplotlib and insert as image."""
        labels = chart_spec.get("labels", [])
        values = chart_spec.get("values", [])
        chart_type = chart_spec.get("type", "bar")

        if not labels or not values:
            return

        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            fig, ax = plt.subplots(figsize=(6, 3.5))

            if chart_type == "bar":
                ax.bar(labels, values, color="#4472C4")
            elif chart_type == "line":
                ax.plot(labels, values, marker="o", color="#4472C4")
            elif chart_type == "pie":
                ax.pie(values, labels=labels, autopct="%1.1f%%")

            plt.tight_layout()
            chart_path = output_dir / f"chart_{uuid.uuid4().hex[:8]}.png"
            fig.savefig(str(chart_path), dpi=150, bbox_inches="tight")
            plt.close(fig)

            doc.add_picture(str(chart_path), width=Inches(5.5))
            chart_path.unlink(missing_ok=True)

        except Exception as e:
            logger.warn("docx_generator.chart_failed", error=str(e))
            doc.add_paragraph(f"[Chart could not be rendered: {e}]")


def _sanitize(text: str) -> str:
    return "".join(c if c.isalnum() or c in " -_" else "" for c in text).strip()[:50]
