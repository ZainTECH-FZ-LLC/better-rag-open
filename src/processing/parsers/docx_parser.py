"""DOCX parser — python-docx for heading hierarchy, styles, and structure."""

from __future__ import annotations

import asyncio
import io

import structlog

from src.processing.parsers.base import DocumentParser, ParsedDocument

logger = structlog.get_logger()


class DOCXParser(DocumentParser):
    async def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, file_bytes, filename)

    def _parse_sync(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))

        text_parts = []
        sections = []
        tables = []

        # Extract paragraphs with heading hierarchy
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            text_parts.append(text)

            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                sections.append({
                    "heading": text,
                    "level": level,
                })

        # Extract tables
        for table in doc.tables:
            table_data = {
                "rows": [],
                "headers": [],
            }
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    table_data["headers"] = cells
                table_data["rows"].append(cells)
            tables.append(table_data)

        full_text = "\n\n".join(text_parts)

        # Extract core properties
        props = {}
        if doc.core_properties:
            cp = doc.core_properties
            props = {
                "author": cp.author,
                "title": cp.title,
                "subject": cp.subject,
                "keywords": cp.keywords,
                "created": str(cp.created) if cp.created else None,
                "modified": str(cp.modified) if cp.modified else None,
            }

        logger.info(
            "docx_parser.parsed",
            filename=filename,
            paragraphs=len(text_parts),
            sections=len(sections),
            tables=len(tables),
        )

        return ParsedDocument(
            text=full_text,
            sections=sections,
            tables=tables,
            word_count=len(full_text.split()),
            file_properties=props,
        )
